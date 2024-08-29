import os
import requests
import csv
import io
import telebot
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton
from docx import Document
from bot_credentials import TELEGRAM_BOT_TOKEN, GOOGLE_BOOKS_API_KEY

bot = telebot.TeleBot(TELEGRAM_BOT_TOKEN)

READING_LIST_FILE = 'reading_list.docx'
user_data = {}

@bot.message_handler(commands=['start'])
def command_start(message):
    bot.reply_to(message, """Welcome to PagePal! 
    I can help you find and manage book recommendations.
    Use /help to see available commands.""")

@bot.message_handler(commands=['help'])
def command_help(message):
    help_text = """
    Available commands:
    /start - Get a welcome message
    /book - Search for books by genre
    /preview - Get a preview link for a book
    /list - Manage your reading list
    /reading_list - View your reading list
    /help - Show this help message
    """
    bot.reply_to(message, help_text)

@bot.message_handler(commands=['book'])
def command_book(message):
    bot.reply_to(message, "Please enter the genre of books you're interested in:")
    user_data[message.chat.id] = {'expect_genre': True}

@bot.message_handler(func=lambda message: user_data.get(message.chat.id, {}).get('expect_genre'))
def handle_genre(message):
    genre = message.text
    base_url = "https://www.googleapis.com/books/v1/volumes"
    params = {
        'q': f'subject:{genre}', 
        'maxResults': 10,  
        'printType': 'books',  
        'key': GOOGLE_BOOKS_API_KEY
    }

    response = requests.get(base_url, params=params)

    if response.status_code == 200 and 'items' in response.json():
        books = response.json()['items']

        
        with io.StringIO() as csv_file:
            writer = csv.writer(csv_file)
            writer.writerow(['Title', 'Author', 'Description', 'Year Published', 'Language', 'Preview Link'])

            for book in books:
                volume_info = book['volumeInfo']
                title = volume_info.get('title', 'N/A')
                authors = ', '.join(volume_info.get('authors', ['N/A']))
                description = volume_info.get('description', 'N/A')[:100] + '...' if volume_info.get('description') else 'N/A'
                published_date = volume_info.get('publishedDate', 'N/A')[:4]
                language = volume_info.get('language', 'N/A')
                preview_link = volume_info.get('previewLink', 'N/A')

                writer.writerow([title, authors, description, published_date, language, preview_link])

            csv_content = csv_file.getvalue()

        with io.BytesIO(csv_content.encode()) as file:
            file.name = f'{genre}_books.csv'
            bot.send_document(message.chat.id, file, caption=f"Books for genre: {genre}")
    else:
        bot.reply_to(message, 'No books found for the given genre.')

    user_data[message.chat.id]['expect_genre'] = False

@bot.message_handler(commands=['preview'])
def preview_command(message):
    bot.reply_to(message, "Please enter the name of the book you want a preview for:")
    user_data[message.chat.id] = {'expect_book_name': True}

@bot.message_handler(func=lambda message: user_data.get(message.chat.id, {}).get('expect_book_name'))
def handle_book_name(message):
    book_name = message.text
    base_url = "https://www.googleapis.com/books/v1/volumes"
    params = {
        'q': book_name,
        'maxResults': 1,
        'printType': 'books',
        'key': GOOGLE_BOOKS_API_KEY
    }

    response = requests.get(base_url, params=params)

    if response.status_code == 200 and 'items' in response.json():
        preview_link = response.json()['items'][0]['volumeInfo'].get('previewLink', 'No preview available')
        bot.reply_to(message, f'Preview link for "{book_name}": {preview_link}')
    else:
        bot.reply_to(message, 'No preview found for the given book.')

    user_data[message.chat.id]['expect_book_name'] = False

@bot.message_handler(commands=['list'])
def list_command(message):
    bot.reply_to(message, 'Please enter the name of the book you want to add to your reading list:')
    user_data[message.chat.id] = {'expect_list_book': True}

@bot.message_handler(func=lambda message: user_data.get(message.chat.id, {}).get('expect_list_book'))
def handle_list_book(message):
    book_name = message.text
    user_data[message.chat.id]['list_book_name'] = book_name
    bot.reply_to(message, 'Now, please use the /reading_list command to manage your reading list.')
    user_data[message.chat.id]['expect_list_book'] = False

@bot.message_handler(commands=['reading_list'])
def reading_list_command(message):
    markup = InlineKeyboardMarkup()
    markup.row_width = 1
    markup.add(
        InlineKeyboardButton("Add a book", callback_data='add_book'),
        InlineKeyboardButton("Delete a book", callback_data='delete_book'),
        InlineKeyboardButton("View Reading List", callback_data='view_list')
    )
    bot.send_message(message.chat.id, "Reading List Management:", reply_markup=markup)

@bot.callback_query_handler(func=lambda call: True)
def button_callback(call):
    if call.data == 'add_book':
        book_name = user_data.get(call.message.chat.id, {}).get('list_book_name')
        if book_name:
            add_to_reading_list(book_name)
            bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
            text=f'Added "{book_name}" to your reading list.')
        else:
            bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
            text='Please use /list command first to specify a book.')

    elif call.data == 'delete_book':
        bot.edit_message_text(chat_id=call.message.chat.id, message_id=call.message.message_id,
            text='working on it.')

    elif call.data == 'view_list':
        reading_list = view_reading_list()
        with io.BytesIO(reading_list) as file:
            file.name = READING_LIST_FILE
            bot.send_document(call.message.chat.id, file, caption="Your Reading List")
def add_to_reading_list(book_name):
    base_url = "https://www.googleapis.com/books/v1/volumes"
    params = {
        'q': book_name,
        'maxResults': 1,
        'printType': 'books',
        'key': GOOGLE_BOOKS_API_KEY
    }
    doc.add_heading('My Reading List', 0)
    response = requests.get(base_url, params=params)
    if response.status_code == 200 and 'items' in response.json():
        book_info = response.json()['items'][0]['volumeInfo']
        title = book_info.get('title', 'N/A')
        preview_link = book_info.get('previewLink', 'N/A')

        doc = Document()
        
        if os.path.exists(READING_LIST_FILE):
            existing_doc = Document(READING_LIST_FILE)
            for paragraph in existing_doc.paragraphs:
                doc.add_paragraph(paragraph.text)

        doc.add_paragraph(f'{title}\n - Preview: {preview_link}')
        doc.save(READING_LIST_FILE)

def view_reading_list():
    if os.path.exists(READING_LIST_FILE):
        with open(READING_LIST_FILE, 'rb') as file:
            return file.read()
    else:
        doc = Document()
        doc.add_heading('My Reading List', 0)
        doc.add_paragraph('Your reading list is empty.')
        bio = io.BytesIO()
        doc.save(bio)
        bio.seek(0)
        return bio.read()

if __name__ == "__main__":
    bot.polling()
# to fix : delete_book, time out error, my reading list
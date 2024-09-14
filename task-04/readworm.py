import os
import requests
import csv
import io
import telebot
from telebot.types import InlineKeyboardMarkup, InlineKeyboardButton
from docx import Document
from docx.shared import Pt
from bot_credentials import TELEGRAM_BOT_TOKEN, GOOGLE_BOOKS_API_KEY

bot = telebot.TeleBot(TELEGRAM_BOT_TOKEN)

READING_LIST_FILE = 'reading_list.docx'

def get_book_info(book_name):
    url = "https://www.googleapis.com/books/v1/volumes?q={}&key={}".format(book_name, GOOGLE_BOOKS_API_KEY)
    response = requests.get(url)
    if response.status_code == 200 and 'items' in response.json():
        book = response.json()['items'][0]['volumeInfo']
        return book
    return None

def get_reading_list():
    if os.path.exists(READING_LIST_FILE):
        doc = Document(READING_LIST_FILE)
        return [para.text for para in doc.paragraphs[1:]]
    return []

def save_reading_list(reading_list):
    doc = Document()
    title = doc.add_heading('My Reading List', 0)
    
    for book in reading_list:
        parts = book.split('\n')
        title = parts[0]
        preview = parts[1] if len(parts) > 1 else ''
        
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True

        if preview:
            p = doc.add_paragraph()
            p.add_run(preview)
    
    doc.save(READING_LIST_FILE)
    
@bot.message_handler(commands=['start'])
def send_welcome(message):
    bot.reply_to(message, "Welcome to PagePal! I can help you find and manage book recommendations. Use /help to see available commands.")

@bot.message_handler(commands=['help'])
def send_help(message):
    help_text = """
    Available commands:
    /start - Get a welcome message
    /book - Search for books by genre
    /preview - Get a preview link for a book
    /list - Add a book to your reading list
    /reading_list - Manage your reading list
    /help - Show this help message
    """
    bot.reply_to(message, help_text)

@bot.message_handler(commands=['book'])
def search_books(message):
    bot.reply_to(message, "Please enter the genre of books you're interested in:")
    bot.register_next_step_handler(message, handle_genre)

def handle_genre(message):
    genre = message.text
    url = "https://www.googleapis.com/books/v1/volumes?q=subject:{}&key={}".format(genre, GOOGLE_BOOKS_API_KEY)
    response = requests.get(url)
    books = response.json().get('items', [])

    if books:
        csv_content = io.StringIO()
        writer = csv.writer(csv_content)
        writer.writerow(['Title', 'Author', 'Description', 'Year', 'Language', 'Preview Link'])

        for book in books[:5]:
            info = book['volumeInfo']
            writer.writerow([
                info.get('title', 'N/A'),
                ', '.join(info.get('authors', ['N/A'])),
                (info.get('description', 'N/A')[:100] + '...') if info.get('description') else 'N/A',
                info.get('publishedDate', 'N/A')[:4],
                info.get('language', 'N/A'),
                info.get('previewLink', 'N/A')
            ])

        csv_file = io.BytesIO(csv_content.getvalue().encode())
        csv_file.name = '{}_books.csv'.format(genre)
        bot.send_document(message.chat.id, csv_file, caption="Books for genre: {}".format(genre))
    else:
        bot.reply_to(message, "No books found for that genre.")

@bot.message_handler(commands=['preview'])
def get_preview(message):
    bot.reply_to(message, "Please enter the name of the book you want a preview for:")
    bot.register_next_step_handler(message, send_preview_link)

def send_preview_link(message):
    book_name = message.text
    book_info = get_book_info(book_name)
    if book_info:
        preview_link = book_info.get('previewLink', 'No preview available')
        bot.reply_to(message, 'Preview link for "{}": {}'.format(book_name, preview_link))
    else:
        bot.reply_to(message, 'No preview found for the given book.')

@bot.message_handler(commands=['list'])
def add_to_list(message):
    bot.reply_to(message, 'Please enter the name of the book you want to add to your reading list:')
    bot.register_next_step_handler(message, process_add_book)

def process_add_book(message):
    book_name = message.text
    book_info = get_book_info(book_name)
    if book_info:
        title = book_info.get('title', 'N/A')
        preview_link = book_info.get('previewLink', 'N/A')
        reading_list = get_reading_list()
        new_entry = '{}\n - Preview: {}'.format(title, preview_link)
        if new_entry in reading_list:
            bot.reply_to(message, 'This book is already in your reading list.')
        else:
            reading_list.append(new_entry)
            save_reading_list(reading_list)
            bot.reply_to(message, 'Added "{}" to your reading list. Use /reading_list to manage your list.'.format(title))
    else:
        bot.reply_to(message, 'Could not find information for the given book.')

@bot.message_handler(commands=['reading_list'])
def manage_reading_list(message):
    markup = InlineKeyboardMarkup()
    markup.row_width = 1
    markup.add(
        InlineKeyboardButton("Add a book", callback_data='add_book'),
        InlineKeyboardButton("Delete a book", callback_data='delete_book'),
        InlineKeyboardButton("View Reading List", callback_data='view_list'),
        InlineKeyboardButton("Clear Reading List", callback_data='clear_list')
    )
    bot.send_message(message.chat.id, "Reading List Management:", reply_markup=markup)

@bot.callback_query_handler(func=lambda call: True)
def callback_query(call):
    if call.data == 'add_book':
        bot.answer_callback_query(call.id)
        bot.send_message(call.message.chat.id, "Please use the /list command to add a book.")

    elif call.data == 'delete_book':
        reading_list = get_reading_list()
        if reading_list:
            markup = InlineKeyboardMarkup()
            for i, book in enumerate(reading_list):
                markup.add(InlineKeyboardButton(book.split('\n')[0], callback_data='delete_{}'.format(i)))
            bot.edit_message_text("Select a book to delete:", call.message.chat.id, call.message.message_id, reply_markup=markup)
        else:
            bot.edit_message_text("Your reading list is empty.", call.message.chat.id, call.message.message_id)

    elif call.data.startswith('delete_'):
        index = int(call.data.split('_')[1])
        reading_list = get_reading_list()
        if 0 <= index < len(reading_list):
            deleted_book = reading_list.pop(index)
            save_reading_list(reading_list)
            bot.answer_callback_query(call.id, "Deleted {} from your reading list.".format(deleted_book.split('\n')[0]))
            bot.edit_message_text("Book deleted from your reading list.", call.message.chat.id, call.message.message_id)
        else:
            bot.answer_callback_query(call.id, "Invalid book selection.")

    elif call.data == 'view_list':
        if os.path.exists(READING_LIST_FILE):
            with open(READING_LIST_FILE, 'rb') as doc:
                bot.send_document(call.message.chat.id, doc, caption="Your Reading List")
        else:
            bot.send_message(call.message.chat.id, "Your reading list is empty.")

    elif call.data == 'clear_list':
        if os.path.exists(READING_LIST_FILE):
            os.remove(READING_LIST_FILE)
            bot.answer_callback_query(call.id, "Your reading list has been cleared.")
            bot.edit_message_text("Your reading list has been cleared.", call.message.chat.id, call.message.message_id)
        else:
            bot.answer_callback_query(call.id, "Your reading list is already empty.")

# Start the bot
bot.polling()